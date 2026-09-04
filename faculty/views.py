import logging
import json
import os
import tempfile
import urllib.error
import urllib.request
from io import BytesIO

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import OperationalError, ProgrammingError
from django.db import transaction
from django.http import HttpResponse
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import AllowAny
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework_simplejwt.tokens import RefreshToken
from api.permissions import IsFacultyUser

from admins.models import AdminInstitution
from backend.services.ollama_service import (
    OllamaConnectionError,
    OllamaInvalidResponseError,
    OllamaModelNotInstalledError,
    OllamaService,
    OllamaServiceError,
    OllamaTimeoutError,
)
from backend.services.document_extractor import extract_document_text
from .models import FacultyProfile, Paper, PaperQuestion, SyllabusUpload
from .serializers import (
    GeneratePaperSerializer,
    FacultyLoginSerializer,
    PaperHistorySerializer,
    PaperSerializer,
    FacultyProfileDepartmentUpdateSerializer,
    FacultyProfileSerializer,
    FacultySignupSerializer,
    SyllabusUploadSerializer,
)
logger = logging.getLogger(__name__)

OLLAMA_UI_MODEL = 'ollama-qwen2.5-3b'
VALID_QUESTION_TYPES = {'Multiple Choice', 'True/False', 'Subjective', 'Fill in the Blanks'}
VALID_DIFFICULTIES = {'Easy', 'Medium', 'Hard'}

QUESTION_TYPE_LABELS = {
    'MCQ': 'MCQ',
    'Subjective': 'Subjective',
    'True/False': 'True/False',
    'Fill in the Blanks': 'Fill in the Blanks',
}


def _error_response(exc, log_message, fallback_message):
    logger.exception(log_message)
    message = str(exc).strip()
    if settings.DEBUG and message:
        response_message = message
    else:
        response_message = fallback_message

    return Response(
        {
            'status': 'error',
            'message': response_message,
        },
        status=status.HTTP_500_INTERNAL_SERVER_ERROR,
    )


def _build_token_payload(user):
    refresh = RefreshToken.for_user(user)
    refresh['role'] = 'faculty'
    access = refresh.access_token
    access['role'] = 'faculty'
    return {
        'refresh': str(refresh),
        'access': str(access),
    }


def _get_or_create_faculty_profile(user):
    profile, _ = FacultyProfile.objects.get_or_create(user=user)
    if not profile.institution:
        default_institution = AdminInstitution.objects.order_by('id').first()
        if default_institution:
            profile.institution = default_institution
            profile.save(update_fields=['institution'])
    profile.ensure_employee_id()
    return profile


def _faculty_or_error(request):
    profile = _get_or_create_faculty_profile(request.user)
    if not profile.institution:
        return None, Response(
            {'status': 'error', 'message': 'Faculty institution not found.'},
            status=status.HTTP_400_BAD_REQUEST,
        )
    return profile, None


def _validate_syllabus_file(uploaded_file):
    filename = uploaded_file.name or ''
    extension = os.path.splitext(filename)[1].lower()
    if extension not in ('.pdf', '.docx'):
        return 'Only PDF and DOCX syllabus files are supported.'
    if uploaded_file.size > 10 * 1024 * 1024:
        return 'Syllabus file must be 10MB or smaller.'
    return None


def _extract_pdf_text(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return '\n'.join(page.extract_text() or '' for page in pdf.pages).strip()
    except ImportError:
        try:
            import fitz
            document = fitz.open(path)
            return '\n'.join(page.get_text() for page in document).strip()
        except ImportError as exc:
            raise ValueError('PDF extraction requires pdfplumber or PyMuPDF to be installed.') from exc


def _extract_docx_text(path):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError('DOCX extraction requires python-docx to be installed.') from exc
    document = Document(path)
    return '\n'.join(paragraph.text for paragraph in document.paragraphs).strip()


def _extract_syllabus_text(uploaded_file):
    return extract_document_text(uploaded_file)


def _paper_prompt(validated_data):
    question_plan = _build_question_plan(validated_data)
    question_plan_text = '\n'.join(
        f"- Question {item['id']}: type={item['type']}, difficulty={item['difficulty']}, marks={item['marks']}"
        for item in question_plan
    )
    return (
        'You are generating an academic examination paper strictly from the uploaded source material and selected topics. '
        'Return ONLY valid JSON with keys: exam_title, duration, total_marks, questions. '
        'The JSON MUST contain a non-empty questions array. Do not return markdown, code fences, commentary, or extra text.\n\n'
        f"Exam title: {validated_data['title']}\n"
        f"Duration: {validated_data['duration']}\n"
        f"Total marks: {validated_data['total_marks']}\n"
        f"Topics: {', '.join(validated_data['topics'])}\n"
        f"Question types: {', '.join(validated_data['question_types'])}\n"
        f"Easy: {validated_data['difficulty_distribution']['Easy']}%\n"
        f"Medium: {validated_data['difficulty_distribution']['Medium']}%\n"
        f"Hard: {validated_data['difficulty_distribution']['Hard']}%\n\n"
        'Generate exactly these questions in the same order and use the same type, difficulty, and marks for each item. '
        'Do not copy this plan into the output; it is only guidance for generation.\n'
        f"{question_plan_text}\n\n"
        'Each question object must use this schema exactly: '
        '{"id": 1, "type": "MCQ", "difficulty": "Easy", "question": "", "options": ["", "", "", ""], "answer": "", "marks": 30}. '
        'For MCQ, use exactly 4 options and set answer to the exact option text. '
        'For True/False, use a meaningful statement and set answer to True or False. '
        'For Fill in the Blanks, use ___ in the question and provide the exact answer. '
        'For Subjective, provide a 2-4 sentence model answer in answer.\n\n'
        'Rules:\n'
        '- Use only the syllabus content and selected topics.\n'
        '- Multiple Choice questions must have exactly 4 options and exactly 1 correct answer.\n'
        '- True/False questions must be meaningful statements with a clear true or false answer.\n'
        '- Fill in the Blanks questions must use ___.\n'
        '- Subjective questions must require reasoning and include a 2-4 sentence model answer.\n'
        '- Total marks of all questions must equal Total Marks exactly.\n'
        '- Follow the requested difficulty distribution as closely as mathematically possible.\n\n'
        'Use this syllabus text:\n'
        f"{validated_data['syllabus_upload'].extracted_text[:18000]}"
    )


def _correction_prompt(validated_data, issues):
    return (
        'The previous JSON was invalid for the exam paper requirements. '
        'Regenerate the entire paper as valid JSON only. No markdown or extra text. '
        'Make sure every question includes question_text, options, correct_answer, and marks.\n\n'
        f"Validation issues: {issues}\n\n"
        f"{_paper_prompt(validated_data)}"
    )


def _build_question_plan(validated_data):
    question_types = list(validated_data['question_types']) or ['MCQ']
    marks_by_difficulty = {}
    allocated_marks = 0
    for difficulty in ('Easy', 'Medium'):
        marks = round(validated_data['total_marks'] * validated_data['difficulty_distribution'][difficulty] / 100)
        marks_by_difficulty[difficulty] = marks
        allocated_marks += marks
    marks_by_difficulty['Hard'] = validated_data['total_marks'] - allocated_marks

    plan = []
    for index, difficulty in enumerate(('Easy', 'Medium', 'Hard')):
        marks = marks_by_difficulty[difficulty]
        if marks <= 0:
            continue
        plan.append({
            'id': index + 1,
            'type': question_types[index % len(question_types)],
            'difficulty': difficulty,
            'marks': marks,
        })
    return plan


def _sentence_count(text):
    fragments = [fragment.strip() for fragment in text.replace('?', '.').replace('!', '.').split('.')]
    return sum(1 for fragment in fragments if fragment)


def _normalize_question_type(question_type):
    mapping = {
        'mcq': 'MCQ',
        'mcqs': 'MCQ',
        'multiple choice': 'MCQ',
        'multiple_choice': 'MCQ',
        'true/false': 'True/False',
        'true false': 'True/False',
        'true_false': 'True/False',
        'subjective': 'Subjective',
        'subjective question': 'Subjective',
        'long answer': 'Subjective',
        'fill in blank': 'Fill in the Blanks',
        'fill in the blank': 'Fill in the Blanks',
        'fill in the blanks': 'Fill in the Blanks',
        'fill in blanks': 'Fill in the Blanks',
    }
    cleaned = str(question_type or '').strip()
    return mapping.get(cleaned.lower(), cleaned)


def _clean_option_text(option):
    cleaned = str(option or '').strip()
    if len(cleaned) > 2 and cleaned[0].isalpha() and cleaned[1] in {')', '.', ':', '-'}:
        cleaned = cleaned[2:].strip()
    if len(cleaned) > 2 and cleaned[0].isdigit() and cleaned[1] in {')', '.', ':', '-'}:
        cleaned = cleaned[2:].strip()
    return cleaned


def _normalize_options(question):
    options = question.get('options')
    if options is None:
        options = question.get('choices', [])
    if isinstance(options, dict):
        options = [options[key] for key in ('A', 'B', 'C', 'D') if options.get(key) is not None]
    if not isinstance(options, list):
        return []
    return [_clean_option_text(option) for option in options]


def _validate_generated_paper(generated, validated_data):
    if not isinstance(generated, dict):
        raise ValueError('AI provider returned invalid JSON.')

    questions = generated.get('questions')
    if not isinstance(questions, list) or not questions:
        raise ValueError('AI provider did not return any questions.')

    question_plan = _build_question_plan(validated_data)
    if len(questions) != len(question_plan):
        raise ValueError('AI provider returned the wrong number of questions.')

    allowed_types = set(validated_data['question_types'])
    total_marks = 0
    marks_by_difficulty = {'Easy': 0, 'Medium': 0, 'Hard': 0}
    normalized_questions = []
    seen_questions = set()

    for index, question in enumerate(questions, start=1):
        if not isinstance(question, dict):
            raise ValueError('AI generated invalid question data.')

        planned_question = question_plan[index - 1]

        question_type = _normalize_question_type(question.get('type') or planned_question['type'])
        difficulty = str(question.get('difficulty') or planned_question['difficulty']).strip().title()
        marks = int(question.get('marks') or planned_question['marks'] or 0)
        text = str(question.get('question') or question.get('question_text') or question.get('text') or '').strip()
        answer = str(
            question.get('answer')
            or question.get('correct_answer')
            or question.get('correct_option')
            or question.get('solution')
            or ''
        ).strip()
        options = _normalize_options(question)

        if question_type not in allowed_types:
            raise ValueError('AI generated a question type that was not selected.')
        if difficulty not in VALID_DIFFICULTIES:
            raise ValueError('AI generated an invalid difficulty level.')
        if marks <= 0:
            raise ValueError('AI generated invalid question marks.')
        if not text:
            raise ValueError('AI generated an invalid question.')
        normalized_key = text.casefold()
        if normalized_key in seen_questions:
            raise ValueError('AI generated duplicate questions.')
        seen_questions.add(normalized_key)

        if question_type == 'MCQ':
            if len(options) != 4:
                raise ValueError('MCQ questions must include exactly 4 options.')
            cleaned_options = options
            if any(not option for option in cleaned_options):
                raise ValueError('MCQ questions must include exactly 4 options.')
            cleaned_answer = _clean_option_text(answer)
            if cleaned_answer in {'A', 'B', 'C', 'D'}:
                cleaned_answer = cleaned_options[ord(cleaned_answer.upper()) - 65]
            if cleaned_answer not in cleaned_options:
                normalized_answer = next((option for option in cleaned_options if option.casefold() == cleaned_answer.casefold()), '')
                cleaned_answer = normalized_answer
            if not cleaned_answer:
                raise ValueError('MCQ questions must include exactly one correct answer.')
            answer = cleaned_answer
        else:
            cleaned_options = []

        if question_type == 'Fill in the Blanks' and '___' not in text:
            raise ValueError('Fill in the Blanks questions must use ___.')
        if question_type == 'True/False' and answer.lower() not in {'true', 'false'}:
            raise ValueError('True/False questions must have a true or false answer.')
        if question_type == 'Subjective':
            if not answer:
                raise ValueError('Subjective questions must include a model answer.')
            if not 2 <= _sentence_count(answer) <= 4:
                raise ValueError('Subjective model answers must be 2 to 4 sentences long.')

        total_marks += marks
        marks_by_difficulty[difficulty] += marks
        normalized_questions.append({
            'question_number': index,
            'type': question_type,
            'difficulty': difficulty,
            'marks': marks,
            'question': text,
            'options': cleaned_options,
            'answer': answer,
        })

    if total_marks != validated_data['total_marks']:
        raise ValueError('Generated question marks do not equal Total Marks.')

    expected_marks = {}
    allocated_marks = 0
    for difficulty in ('Easy', 'Medium'):
        marks = round(validated_data['total_marks'] * validated_data['difficulty_distribution'][difficulty] / 100)
        expected_marks[difficulty] = marks
        allocated_marks += marks
    expected_marks['Hard'] = validated_data['total_marks'] - allocated_marks

    if marks_by_difficulty != expected_marks:
        raise ValueError('Generated question difficulty marks do not match the requested distribution.')

    generated['questions'] = normalized_questions
    generated['exam_title'] = validated_data['title']
    generated['duration'] = validated_data['duration']
    generated['total_marks'] = validated_data['total_marks']
    return generated


def _generate_with_ollama(validated_data):
    service = OllamaService()
    prompt = _paper_prompt(validated_data)
    try:
        generated = service.generate(OLLAMA_UI_MODEL, prompt)
        return _validate_generated_paper(generated, validated_data)
    except (OllamaConnectionError, OllamaTimeoutError, OllamaModelNotInstalledError) as exc:
        raise exc
    except (OllamaInvalidResponseError, OllamaServiceError, ValueError) as exc:
        corrected = service.generate(OLLAMA_UI_MODEL, _correction_prompt(validated_data, str(exc)))
        try:
            return _validate_generated_paper(corrected, validated_data)
        except (OllamaInvalidResponseError, OllamaServiceError, ValueError) as second_exc:
            raise OllamaInvalidResponseError(
                f'Ollama returned an unusable exam paper. {second_exc}'
            ) from second_exc


def _save_generated_paper(faculty, validated_data, generated):
    upload = validated_data['syllabus_upload']
    with transaction.atomic():
        upload.original_file.open('rb')
        paper = Paper.objects.create(
            faculty=faculty,
            institution=faculty.institution,
            title=validated_data['title'],
            syllabus_filename=upload.original_filename,
            extracted_syllabus_text=upload.extracted_text,
            ai_model=validated_data['model'],
            topics=validated_data['topics'],
            question_types=validated_data['question_types'],
            difficulty_distribution=validated_data['difficulty_distribution'],
            duration=validated_data['duration'],
            total_marks=validated_data['total_marks'],
            generated_questions=generated['questions'],
            is_published=True,
        )
        paper.syllabus_file.save(upload.original_filename, ContentFile(upload.original_file.read()), save=True)

        PaperQuestion.objects.bulk_create([
            PaperQuestion(
                paper=paper,
                question_number=question['question_number'],
                question_type=question['type'],
                difficulty=question['difficulty'],
                marks=question['marks'],
                question=question['question'],
                options=question['options'],
                answer=question['answer'],
            )
            for question in generated['questions']
        ])
    return paper


def _paper_for_faculty(faculty, paper_id):
    return (
        Paper.objects.select_related('institution', 'faculty')
        .prefetch_related('questions')
        .filter(id=paper_id, institution=faculty.institution)
        .first()
    )


def _render_paper_text(paper):
    sections = {'MCQ': [], 'Subjective': [], 'True/False': [], 'Fill in the Blanks': []}
    for question in paper.questions.all():
        sections.setdefault(question.question_type, []).append(question)
    lines = [
        paper.institution.institution_name,
        paper.title,
        f'Duration: {paper.duration} minutes',
        f'Total Marks: {paper.total_marks}',
        '',
        'Instructions: Answer all questions. Marks are shown beside each question.',
        '',
    ]
    section_name = 'A'
    question_number = 1
    for question_type, questions in sections.items():
        if not questions:
            continue
        lines.append(f'Section {section_name} - {question_type}')
        for question in questions:
            lines.append(f'{question_number}. {question.question} [{question.marks} marks]')
            for option in question.options:
                lines.append(f'   - {option}')
            question_number += 1
        lines.append('')
        section_name = chr(ord(section_name) + 1)
    return '\n'.join(lines)


def _export_pdf(paper):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ValueError('PDF export requires reportlab to be installed.') from exc

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    for line in _render_paper_text(paper).splitlines():
        if y < 50:
            pdf.showPage()
            y = height - 50
        pdf.drawString(50, y, line[:110])
        y -= 18
    pdf.save()
    buffer.seek(0)
    return buffer.getvalue(), 'application/pdf', f'{paper.title}.pdf'


def _export_docx(paper):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError('DOCX export requires python-docx to be installed.') from exc
    document = Document()
    document.add_heading(paper.institution.institution_name, level=1)
    document.add_heading(paper.title, level=2)
    document.add_paragraph(f'Duration: {paper.duration} minutes')
    document.add_paragraph(f'Total Marks: {paper.total_marks}')
    document.add_heading('Instructions', level=2)
    document.add_paragraph('Answer all questions. Marks are shown beside each question.')
    for line in _render_paper_text(paper).splitlines()[7:]:
        if line.startswith('Section '):
            document.add_heading(line, level=2)
        elif line:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return (
        buffer.getvalue(),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        f'{paper.title}.docx',
    )


class OllamaStatusAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def get(self, request):
        service = OllamaService()
        status_data = service.get_status(OLLAMA_UI_MODEL, timeout=5)
        return Response(
            {
                'status': 'success',
                'connected': status_data['connected'],
                'model_installed': status_data['model_installed'],
                'message': status_data['message'],
            },
            status=status.HTTP_200_OK,
        )


class FacultySignupAPIView(APIView):
    authentication_classes = []
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = FacultySignupSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        user = serializer.save()

        return Response(
            {
                'status': 'success',
                'message': 'Faculty account created successfully.',
                'tokens': _build_token_payload(user),
                'user': {
                    'id': user.id,
                    'first_name': user.first_name,
                    'last_name': user.last_name,
                    'email': user.email,
                    'role': 'faculty',
                },
            },
            status=status.HTTP_201_CREATED,
        )


class FacultyLoginAPIView(APIView):
    authentication_classes = []
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = FacultyLoginSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        user = serializer.validated_data['user']

        return Response(
            {
                'status': 'success',
                'message': 'Faculty login successful.',
                'tokens': _build_token_payload(user),
                'user': {
                    'id': user.id,
                    'first_name': user.first_name,
                    'last_name': user.last_name,
                    'email': user.email,
                    'role': 'faculty',
                },
            },
            status=status.HTTP_200_OK,
        )


class FacultyProfileAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def get(self, request):
        try:
            profile = _get_or_create_faculty_profile(request.user)
        except (OperationalError, ProgrammingError) as exc:
            return _error_response(
                exc,
                f'Faculty profile fetch database error for user_id={request.user.id}',
                'Database schema issue detected. Please run migrations.',
            )
        except Exception as exc:
            return _error_response(
                exc,
                f'Unexpected faculty profile fetch error for user_id={request.user.id}',
                'Unable to fetch faculty profile.',
            )
        serializer = FacultyProfileSerializer(profile)
        return Response(
            {
                'status': 'success',
                'message': 'Faculty profile fetched successfully.',
                'data': serializer.data,
            },
            status=status.HTTP_200_OK,
        )

    def put(self, request):
        try:
            profile = _get_or_create_faculty_profile(request.user)
        except (OperationalError, ProgrammingError) as exc:
            return _error_response(
                exc,
                f'Faculty profile update database error for user_id={request.user.id}',
                'Database schema issue detected. Please run migrations.',
            )
        except Exception as exc:
            return _error_response(
                exc,
                f'Unexpected faculty profile update error for user_id={request.user.id}',
                'Unable to update faculty profile.',
            )

        serializer = FacultyProfileDepartmentUpdateSerializer(profile, data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()

        response_serializer = FacultyProfileSerializer(profile)
        return Response(
            {
                'status': 'success',
                'message': 'Department updated successfully.',
                'data': response_serializer.data,
            },
            status=status.HTTP_200_OK,
        )


class SyllabusUploadAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response

        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return Response(
                {'status': 'error', 'message': 'Syllabus file is required.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        validation_error = _validate_syllabus_file(uploaded_file)
        if validation_error:
            return Response({'status': 'error', 'message': validation_error}, status=status.HTTP_400_BAD_REQUEST)

        try:
            extracted_text = _extract_syllabus_text(uploaded_file)
            uploaded_file.seek(0)
            syllabus = SyllabusUpload.objects.create(
                faculty=faculty,
                institution=faculty.institution,
                original_file=uploaded_file,
                original_filename=uploaded_file.name,
                content_type=getattr(uploaded_file, 'content_type', '') or '',
                extracted_text=extracted_text,
            )
        except ValueError as exc:
            return Response({'status': 'error', 'message': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as exc:
            return _error_response(exc, 'Unexpected syllabus upload error', 'Unable to upload syllabus.')

        return Response(
            {
                'status': 'success',
                'message': 'Syllabus uploaded successfully.',
                'data': SyllabusUploadSerializer(syllabus).data,
            },
            status=status.HTTP_201_CREATED,
        )


class GeneratePaperAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def post(self, request):
        logger.info('[GeneratePaper] POST request received')
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response
        logger.info('[GeneratePaper] User authenticated; faculty identified')

        serializer = GeneratePaperSerializer(data=request.data, context={'faculty': faculty})
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        logger.info('[GeneratePaper] Request data parsed; institution identified')
        upload = validated_data['syllabus_upload']
        logger.info('[GeneratePaper] Syllabus resolved; extracted text length: %s', len(upload.extracted_text or ''))
        if not upload.original_file or not upload.original_file.name or not upload.original_file.storage.exists(upload.original_file.name):
            return Response(
                {
                    'success': False,
                    'provider': None,
                    'error_code': 'syllabus_not_found',
                    'message': 'The uploaded syllabus file is no longer available. Please upload it again.',
                    'fallback_attempted': False,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )
        if not str(upload.extracted_text or '').strip():
            return Response(
                {
                    'success': False,
                    'status': 'error',
                    'provider': 'Ollama',
                    'error_code': 'empty_syllabus',
                    'message': 'Unable to generate paper. Syllabus content is empty.',
                    'fallback_attempted': False,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            logger.info('[GeneratePaper] Model selected: %s', OllamaService().model_for(validated_data['model']))
            logger.info('[GeneratePaper] Calling Ollama')
            generated = _generate_with_ollama(validated_data)
            logger.info('[GeneratePaper] Ollama response received; validation passed; marks validation passed')
            logger.info('[GeneratePaper] Saving paper')
            paper = _save_generated_paper(faculty, validated_data, generated)
            logger.info('[GeneratePaper] Paper saved')
        except OllamaModelNotInstalledError:
            return Response(
                {
                    'success': False,
                    'status': 'error',
                    'provider': 'Ollama',
                    'error_code': 'model_not_found',
                    'message': f'Selected model is not installed. Run: ollama pull {OllamaService().model_for(validated_data["model"]) }',
                    'fallback_attempted': False,
                },
                status=status.HTTP_503_SERVICE_UNAVAILABLE,
            )
        except OllamaConnectionError:
            return Response(
                {
                    'success': False,
                    'status': 'error',
                    'provider': 'Ollama',
                    'error_code': 'connection_error',
                    'message': 'Ollama is not running. Please start Ollama and try again.',
                    'fallback_attempted': False,
                },
                status=status.HTTP_503_SERVICE_UNAVAILABLE,
            )
        except OllamaTimeoutError:
            return Response(
                {
                    'success': False,
                    'status': 'error',
                    'provider': 'Ollama',
                    'error_code': 'timeout',
                    'message': 'Ollama generation timed out. Please try again.',
                    'fallback_attempted': False,
                },
                status=status.HTTP_503_SERVICE_UNAVAILABLE,
            )
        except OllamaInvalidResponseError as exc:
            return Response(
                {
                    'success': False,
                    'status': 'error',
                    'provider': 'Ollama',
                    'error_code': 'invalid_ai_response',
                    'error': str(exc),
                    'message': str(exc),
                },
                status=status.HTTP_502_BAD_GATEWAY,
            )
        except ValueError as exc:
            return Response({'success': False, 'status': 'error', 'message': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        except (BrokenPipeError, ConnectionAbortedError):
            logger.info('Client disconnected while generating a paper; generation response was discarded.')
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as exc:
            return _error_response(exc, 'Unexpected paper generation error', 'Unable to generate paper.')

        paper_data = PaperSerializer(paper).data
        response_data = {
            'success': True,
            'status': 'success',
            'message': 'Paper generated successfully.',
            'provider': 'Ollama',
            'model': OllamaService().model_for(validated_data['model']),
            'exam_title': paper.title,
            'duration': paper.duration,
            'total_marks': paper.total_marks,
            'questions': paper_data['questions'],
            'paper': paper_data,
            'data': paper_data,
        }

        return Response(
            response_data,
            status=status.HTTP_200_OK,
        )


class PaperHistoryAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def get(self, request):
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response

        papers = Paper.objects.filter(institution=faculty.institution, faculty=faculty).prefetch_related('questions')
        return Response(
            {
                'status': 'success',
                'message': 'Paper history fetched successfully.',
                'data': PaperHistorySerializer(papers, many=True).data,
            },
            status=status.HTTP_200_OK,
        )


class PaperDetailAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def get(self, request, paper_id):
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response

        paper = _paper_for_faculty(faculty, paper_id)
        if not paper:
            return Response({'status': 'error', 'message': 'Paper not found.'}, status=status.HTTP_404_NOT_FOUND)
        return Response(
            {
                'status': 'success',
                'message': 'Paper fetched successfully.',
                'data': PaperSerializer(paper).data,
            },
            status=status.HTTP_200_OK,
        )

    def delete(self, request, paper_id):
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response

        paper = _paper_for_faculty(faculty, paper_id)
        if not paper:
            return Response({'status': 'error', 'message': 'Paper not found.'}, status=status.HTTP_404_NOT_FOUND)
        paper.delete()
        return Response(
            {'status': 'success', 'message': 'Paper deleted successfully.'},
            status=status.HTTP_200_OK,
        )


class PaperExportAPIView(APIView):
    permission_classes = [IsAuthenticated, IsFacultyUser]

    def get(self, request, paper_id):
        faculty, error_response = _faculty_or_error(request)
        if error_response:
            return error_response

        paper = _paper_for_faculty(faculty, paper_id)
        if not paper:
            return Response({'status': 'error', 'message': 'Paper not found.'}, status=status.HTTP_404_NOT_FOUND)

        export_format = (request.query_params.get('format') or 'pdf').lower()
        try:
            if export_format == 'docx':
                content, content_type, filename = _export_docx(paper)
            elif export_format == 'pdf':
                content, content_type, filename = _export_pdf(paper)
            else:
                return Response({'status': 'error', 'message': 'Export format must be pdf or docx.'}, status=status.HTTP_400_BAD_REQUEST)
        except ValueError as exc:
            return Response({'status': 'error', 'message': str(exc)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as exc:
            return _error_response(exc, 'Unexpected paper export error', 'Unable to export paper.')

        response = HttpResponse(content, content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
